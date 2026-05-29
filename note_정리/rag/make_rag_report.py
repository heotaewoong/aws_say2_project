"""RAG 구현 보고서 DOCX 생성 스크립트 (초보자용)"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ───────────────────────────── 스타일 유틸 ─────────────────────────────

def set_heading(paragraph, level=1):
    colors = {1: "1F4E79", 2: "2E75B6", 3: "2F5496"}
    sizes  = {1: 22, 2: 16, 3: 13}
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 13))
    run.font.color.rgb = RGBColor.from_string(colors.get(level, "000000"))
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after  = Pt(4)

def h1(doc, text):
    p = doc.add_paragraph(text)
    set_heading(p, 1)
    return p

def h2(doc, text):
    p = doc.add_paragraph(text)
    set_heading(p, 2)
    return p

def h3(doc, text):
    p = doc.add_paragraph(text)
    set_heading(p, 3)
    return p

def body(doc, text, bold_parts=None):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(0.3 * (level + 1))
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0x52)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), "F0F4F8")
    p._p.pPr.append(shading)
    return p

def tip_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("💡  " + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x7B, 0x34, 0x02)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), "FFF3E0")
    p._p.pPr.append(shading)
    return p

def make_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "1F4E79")
        tcPr.append(shd)
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].runs[0].font.size = Pt(9.5)
            if ri % 2 == 0:
                tc = cells[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "EBF5FB")
                tcPr.append(shd)
    if col_widths:
        for row in t.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()
    return t

# ─────────────────────────── 표지 페이지 ───────────────────────────────

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n\n\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Rare-Link AI")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("RAG 파이프라인 구현 보고서")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("희귀 폐질환 진단 보조 AI 시스템")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SKKU AWS SAY 2기 2팀\n권미라 · 배기태 · 허태웅\n확정일: 2026-04-29 | 최종 업데이트: 2026-05-24")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_page_break()

# ─────────────────────────── 목차 ───────────────────────────────────────

h1(doc, "목차")
toc_items = [
    ("1장", "이 프로젝트란?"),
    ("2장", "RAG란 무엇인가 — 쉽게 설명"),
    ("3장", "전체 시스템 아키텍처 (AWS 구성)"),
    ("4장", "RAG 파이프라인 5단계 상세 설명"),
    ("5장", "사용된 외부 데이터 API 5개"),
    ("6장", "AWS 리소스 구성 목록"),
    ("7장", "데이터베이스 구조 (Aurora PostgreSQL)"),
    ("8장", "보안 구성"),
    ("9장", "배포 현황 — CloudFront 사이트"),
    ("10장", "검증 결과"),
    ("11장", "향후 계획"),
    ("부록", "용어 사전"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    run_num = p.add_run(f"{num}  ")
    run_num.bold = True
    run_num.font.size = Pt(11)
    run_num.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run_title = p.add_run(title)
    run_title.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 1장 — 이 프로젝트란?
# ═══════════════════════════════════════════════════════════════════════

h1(doc, "1장. 이 프로젝트란?")
body(doc, "Rare-Link AI는 희귀 폐질환 진단을 도와주는 AI 시스템입니다. 의사가 환자의 정보를 입력하면 AI가 다양한 의학 데이터베이스를 검색하고 진단 보조 소견서를 자동으로 생성해 줍니다.")

h2(doc, "1.1 왜 이 시스템이 필요한가?")
bullet(doc, "희귀 폐질환은 전 세계 수천 가지 이상이며, 의사 혼자 모든 희귀질환 정보를 기억하기 어렵습니다.")
bullet(doc, "진단까지 평균 5~7년이 걸리는 희귀질환 환자들을 더 빨리 도울 수 있습니다.")
bullet(doc, "최신 의학 논문과 임상 데이터를 실시간으로 검색하여 근거 있는 소견을 제공합니다.")

tip_box(doc, "이 시스템은 '최종 진단'을 내리는 게 아니라, 의사의 판단을 돕는 '진단 보조' 도구입니다. 최종 결정은 항상 의사가 내립니다.")

h2(doc, "1.2 어떤 정보를 입력하나요?")
make_table(doc,
    ["입력 정보", "설명", "예시"],
    [
        ["흉부 X-ray", "폐 사진 (JPEG/PNG)", "chest_xray.jpg"],
        ["혈액검사 결과", "혈액 내 수치들", "CRP: 15, WBC: 12000 ..."],
        ["증상 텍스트", "환자가 호소하는 증상", "기침, 호흡곤란, 발열 3일"],
        ["활력징후 (Vital)", "혈압, 맥박, 체온 등", "SpO2: 92%, HR: 110"],
        ["미생물 검사 (Micro)", "균 배양 결과", "BAL culture: no growth"],
        ["환자 기본정보", "나이, 성별, 내원 사유", "64세, 남성, 호흡곤란"],
    ],
    col_widths=[4, 5, 7]
)

h2(doc, "1.3 어떤 결과가 나오나요?")
body(doc, "AI는 다음과 같은 JSON 형태의 진단 보조 소견서를 생성합니다:")
bullet(doc, "즉시 검사 권고 (immediate_workup): 지금 당장 해야 할 검사 목록")
bullet(doc, "전문과 협진 의뢰 (specialist_referral): 어느 과에 의뢰해야 하는지")
bullet(doc, "치료 가이드라인 (treatment_guideline): 각 질환별 치료 방향")
bullet(doc, "유전자 검사 권고 (genetic_test): 필요한 유전자 검사 목록")
bullet(doc, "추가 혈액검사 (additional_lab): 추가로 필요한 검사")
bullet(doc, "임상 소견 요약 (clinical_notes): 진단 근거, 감별진단, RAG 증거 등")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 2장 — RAG란 무엇인가
# ═══════════════════════════════════════════════════════════════════════

h1(doc, "2장. RAG란 무엇인가 — 쉽게 설명")

h2(doc, "2.1 일반 AI vs RAG AI")
body(doc, "일반 AI(예: ChatGPT)는 학습 당시 배운 정보만 사용합니다. 최신 논문이나 특정 환자 데이터는 모릅니다.")
body(doc, "RAG(Retrieval-Augmented Generation)는 답변을 생성하기 전에 외부 데이터베이스를 실시간으로 검색합니다.")

code_block(doc,
"""[일반 AI]
  질문 → AI → 답변 (학습 데이터만 사용, 최신 정보 없음)

[RAG AI]
  질문 → 외부 DB 검색 → 검색 결과 + 질문 → AI → 근거 있는 답변
         (PubMed, Orphanet, ClinicalTrials 등 실시간 검색)"""
)

tip_box(doc, "도서관 비유: 일반 AI는 이미 머릿속에 있는 지식만 쓰는 학생이고, RAG AI는 필요할 때 도서관에서 책을 찾아 읽고 답하는 학생입니다.")

h2(doc, "2.2 이 프로젝트에서 RAG는 어떻게 동작하나요?")
body(doc, "환자 데이터를 분석하여 의심 질환 Top 3를 추린 뒤, 5개의 의학 데이터베이스에서 동시에 관련 정보를 검색합니다. 검색된 정보(RAG 컨텍스트)를 AI에게 제공하면 근거 있는 소견서를 작성합니다.")

h2(doc, "2.3 HPO란?")
body(doc, "HPO(Human Phenotype Ontology)는 사람의 증상을 표준화된 코드로 나타내는 체계입니다.")
bullet(doc, "예: '기침' → HP:0012735")
bullet(doc, "예: '폐의 경화' → HP:0032262")
bullet(doc, "전 세계 의사들이 동일한 언어(코드)로 증상을 표현할 수 있어 데이터베이스 검색이 가능합니다.")
tip_box(doc, "HPO 코드는 의학판 '바코드'라고 생각하면 됩니다. 같은 증상도 나라마다 표현이 다르지만 HPO 코드는 전 세계 공통입니다.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 3장 — 전체 시스템 아키텍처
# ═══════════════════════════════════════════════════════════════════════

h1(doc, "3장. 전체 시스템 아키텍처 (AWS 구성)")

h2(doc, "3.1 사용자 → AWS 서비스 흐름")
code_block(doc,
"""[사용자/의사]
    ↓ 브라우저에서 접속
[CloudFront CDN]  ← 전 세계 빠른 배포 (https://d300v14l8u0wx7.cloudfront.net)
    ↓
[React 프론트엔드 (S3 호스팅)]
    ↓ API 호출
[API Gateway]  ← REST API 엔드포인트
    ↓
[Lambda 함수들]  ← 각 Phase별 처리
    ↓                      ↓
[SageMaker Endpoint]    [EC2 (RAG 실행)]
  SooNet AI 모델          Bedrock + 외부 API 호출
    ↓                      ↓
[S3 버킷]            [Aurora PostgreSQL DB]
  X-ray 저장             환자/세션/소견서 저장"""
)

h2(doc, "3.2 AWS 서비스 역할 한눈에 보기")
make_table(doc,
    ["AWS 서비스", "역할", "리소스 이름"],
    [
        ["CloudFront", "전 세계 빠른 웹 서비스 배포", "say2-2team-cf-distribution"],
        ["S3 버킷", "X-ray 이미지, AI 모델, 프론트엔드 파일 저장", "say2-2team-bucket"],
        ["Lambda", "Phase별 처리 서버리스 함수", "say2-2team-phase2-vision"],
        ["SageMaker", "SooNet AI 모델 추론 (X-ray 분석)", "say2-2team-soonet-endpoint"],
        ["EC2", "RAG 파이프라인 실행 환경", "2-2team-fhir-ec2"],
        ["Bedrock", "Claude AI 모델 호출 (AWS 관리형)", "claude-3-5-sonnet/haiku"],
        ["Aurora PostgreSQL", "환자 데이터, 진단 세션, 소견서 저장 DB", "patient-db-cluster"],
        ["Cognito", "사용자 로그인 인증", "say2-2team-rare-link-pool"],
        ["KMS", "데이터 암호화 키 관리", "alias/say2-2team-data-key"],
        ["CloudTrail", "모든 API 호출 감사 로그", "say2-2team-audit-trail"],
        ["WAF", "웹 방화벽 (DDoS, SQL Injection 차단)", "say2-2team-waf"],
        ["SNS", "이상 탐지 시 이메일 알람", "say2-2team-alerts"],
        ["VPC", "AWS 내부 가상 사설 네트워크", "say2-2team (10.0.0.0/24)"],
    ],
    col_widths=[4, 7, 5.5]
)

h2(doc, "3.3 VPC 네트워크 구조")
body(doc, "보안을 위해 핵심 서비스들은 외부 인터넷과 직접 연결되지 않는 Private Subnet에 배치되어 있습니다.")
make_table(doc,
    ["서브넷", "IP 대역", "배치된 서비스"],
    [
        ["Public Subnet", "10.0.0.0/28", "Internet Gateway, NAT Gateway"],
        ["Private Subnet 1", "10.0.0.128/28", "Lambda, SageMaker Endpoint"],
        ["Private Subnet 2", "10.0.0.16/28", "Phase 4·5 예약"],
        ["Private Subnet 3", "10.0.0.32/28", "RAG & Report 예약"],
    ],
    col_widths=[5, 4, 7.5]
)
tip_box(doc, "Private Subnet = 인터넷에서 직접 접근 불가한 내부 구역. 환자 데이터를 더 안전하게 보호합니다.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 4장 — RAG 파이프라인 5단계
# ═══════════════════════════════════════════════════════════════════════

h1(doc, "4장. RAG 파이프라인 5단계 상세 설명")
body(doc, "환자 데이터가 입력되면 아래 5단계 파이프라인을 순서대로 거쳐 최종 소견서가 생성됩니다.")

code_block(doc,
"""환자 입력 (X-ray + 혈액검사 + 증상 + Vital + Micro)
            │
            ▼
    ① 멀티모달 → HPO 변환  (Phase 1~3)
            │
            ▼
    ② 듀얼 트랙 스코어링   (일반 Top10 + 희귀 리스팅)
            │
            ▼
    ③ Phase 4 LLM 정리     (Haiku → Top 3 통합)
            │
            ▼
    ④ RAG 트리거            (5개 외부 API 병렬 검색)
            │
            ▼
    ⑤ 최종 소견서 생성      (Claude Sonnet 3.5 → JSON)"""
)

# 단계 1
h2(doc, "4.1 단계 ①: 멀티모달 → HPO 변환")
body(doc, "세 가지 서로 다른 정보 소스에서 각각 HPO 코드를 추출합니다.")

h3(doc, "Phase 1 — 증상 텍스트 → HPO (Claude Haiku 사용)")
body(doc, "환자가 표현한 증상 문장을 AWS Bedrock의 Claude Haiku 모델에게 보냅니다. AI가 텍스트를 읽고 해당하는 HPO 코드를 찾아냅니다.")
bullet(doc, "입력: '기침이 심하고 열이 나며 숨이 차다'")
bullet(doc, "출력 Positive HPO: HP:0012735 (기침), HP:0001945 (발열), HP:0002094 (호흡곤란)")
bullet(doc, "출력 Negative HPO: 환자가 '없다'고 말한 증상도 별도로 추출")

h3(doc, "Phase 2 — X-ray → HPO (SooNet AI 모델)")
body(doc, "SooNet은 DenseNet-121 + U-Net 기반의 흉부 X-ray 분석 AI 모델입니다. AWS SageMaker에 배포되어 있으며, 14개 질환 카테고리별 확률을 계산합니다.")
make_table(doc,
    ["분석 항목", "정상 범위 초과 시"],
    [
        ["폐렴 (Pneumonia)", "HP:0002090 추출"],
        ["흉막삼출 (Pleural Effusion)", "HP:0002202 추출"],
        ["폐경화 (Consolidation)", "HP:0032262 추출"],
        ["기흉 (Pneumothorax)", "HP:0002107 추출"],
        ["기타 10개 질환", "각 해당 HPO 코드 추출"],
    ],
    col_widths=[6, 10.5]
)
bullet(doc, "확률 ≥ 0.3 (30%)인 항목만 HPO로 변환")
bullet(doc, "SageMaker Endpoint: say2-2team-soonet-endpoint (ml.m5.large)")

h3(doc, "Phase 3 — 혈액검사/Vital/미생물 → HPO (Rule-based)")
body(doc, "혈액검사 수치가 정상 범위를 벗어나면 자동으로 관련 HPO 코드를 부여합니다. 컴퓨터 규칙으로 처리하므로 빠르고 일관성 있습니다.")
make_table(doc,
    ["검사 항목", "이상 소견", "HPO 코드"],
    [
        ["CRP", "> 1.0 mg/dL", "HP:0011227 (전신 염증)"],
        ["WBC", "> 11,000", "HP:0001974 (백혈구증가증)"],
        ["SpO2", "< 95%", "HP:0012418 (저산소혈증)"],
        ["Ferritin", "상승", "HP:0030736 (과다페리틴혈증)"],
    ],
    col_widths=[4, 4.5, 8]
)

# 단계 2
h2(doc, "4.2 단계 ②: 듀얼 트랙 스코어링")
body(doc, "일반 폐질환과 희귀 폐질환을 분리하여 각각 다른 방식으로 점수를 계산합니다.")

h3(doc, "트랙 A — 일반/기타 폐질환 (Rule-based Top 10)")
body(doc, "폐렴, COPD, 결핵, 폐부종 같은 흔한 폐질환을 대상으로 HPO 매칭 + X-ray 확률 + Lab 수치를 조합하여 점수를 계산합니다. 상위 10개를 추립니다.")

h3(doc, "트랙 B — 희귀 폐질환 (LIRICAL LR 스코어링)")
body(doc, "Orphanet 데이터베이스에 수록된 4,335개의 희귀질환을 대상으로 LIRICAL(Likelihood Ratio-based Inference for Clinical AI Linkage) 알고리즘으로 점수를 계산합니다.")
bullet(doc, "Orphanet CSV: 4,335개 희귀질환 × 115,878행 HPO 빈도 데이터")
bullet(doc, "LR 임계치를 넘은 질환만 '희귀 리스팅'으로 출력 (없으면 빈 목록)")
tip_box(doc, "왜 분리하나요? 일반 질환은 흔하지만 LIRICAL로는 잘 감지 안 되고, 희귀 질환은 Orphanet 빈도 데이터가 있어야 정확합니다. 두 트랙을 각각 돌려야 모두 커버할 수 있습니다.")

# 단계 3
h2(doc, "4.3 단계 ③: Phase 4 — LLM이 Top 3 정리")
body(doc, "일반 Top 10과 희귀 리스팅을 Claude Haiku에게 보냅니다. AI가 임상적 맥락을 고려하여 두 리스트를 통합한 '최우선 Top 3 의심 질환'을 정리합니다.")
bullet(doc, "일반 Top 10의 점수 체계 vs 희귀 리스팅의 LR 점수는 단위가 달라 단순 합산 불가")
bullet(doc, "AI가 임상 근거(증상 + Lab + 희귀 가능성)를 보고 통합 우선순위 결정")
bullet(doc, "출력: [{rank, disease_name, orpha_code, score, source}] × 3개")

# 단계 4
h2(doc, "4.4 단계 ④: RAG 트리거 — 5개 API 병렬 검색")
body(doc, "Top 3 결과를 바탕으로 5개 의학 데이터베이스에 동시에(병렬) 검색 요청을 보냅니다. 동시에 보내서 약 60초 내로 모든 결과를 받습니다.")

h3(doc, "조건부 검색 로직 (어떤 API를 호출할지 결정)")
make_table(doc,
    ["상황", "호출하는 API"],
    [
        ["희귀 리스팅 있음 (케이스 A)", "Orphanet + PubCaseFinder + Monarch + PubMed + ClinicalTrials (5개 전부)"],
        ["리스팅 없음 + Top3에 희귀질환 코드 있음 (케이스 B)", "Orphanet + PubCaseFinder + Monarch + PubMed + ClinicalTrials (5개 전부, 교차검증)"],
        ["Top 3 전부 일반 질환 (케이스 C)", "Monarch + PubMed + ClinicalTrials (3개만, Orphanet/PCF 스킵)"],
    ],
    col_widths=[6.5, 10]
)

# 단계 5
h2(doc, "4.5 단계 ⑤: 최종 소견서 생성")
body(doc, "수집된 모든 정보를 Claude Sonnet 3.5 (AWS Bedrock)에게 전달하여 최종 JSON 소견서를 생성합니다.")
make_table(doc,
    ["설정", "값", "이유"],
    [
        ["AI 모델", "claude-3-5-sonnet-20241022-v2:0 (APAC)", "한국어 + 의학 추론 최고 성능"],
        ["최대 토큰", "2,048", "충분한 소견서 길이 확보"],
        ["Temperature", "0.0", "매번 동일한 결과 (재현성)"],
        ["출력 형식", "JSON only (Markdown 금지)", "프론트엔드 파싱 편의"],
    ],
    col_widths=[4, 5.5, 7]
)
bullet(doc, "MRN(병록번호) 절대 포함 금지 — 개인정보 보호")
bullet(doc, "희귀질환이 Top 3 이내이면 MDT(다학제팀 협진) 권고 필수")
bullet(doc, "모든 주장은 RAG 검색 결과 또는 공인 의학 가이드라인에 근거해야 함")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 5장 — 외부 API 5개
# ═══════════════════════════════════════════════════════════════════════

h1(doc, "5장. 사용된 외부 데이터 API 5개")
body(doc, "이 프로젝트는 5개의 공개 의학 데이터베이스를 실시간으로 활용합니다. 모두 무료(또는 무료 사용 범위 내)입니다.")

make_table(doc,
    ["순서", "API 이름", "제공 기관", "제공 정보", "인증"],
    [
        ["①", "PubCaseFinder", "DBCLS (일본)", "HPO 입력 → 후보 희귀질환 + OMIM 코드 + 관련 논문 PMID", "무료"],
        ["②", "Orphanet (로컬 XML)", "Orphanet (유럽)", "OrphaCode → 관련 유전자, 유병률, 발병 연령, 유전 양식", "무료 XML 다운로드"],
        ["③", "Monarch Initiative", "EMBL-EBI (유럽)", "질환명/코드 → 인과 유전자, HPO 매핑, 질환 설명", "무료"],
        ["④", "PubMed E-utilities", "NCBI (미국)", "질환명 → 최신 케이스 리포트 논문 3건 (PMID + 초록)", "무료 (API Key 권장)"],
        ["⑤", "ClinicalTrials.gov v2", "NIH (미국)", "질환명 → 현재 모집 중인 임상시험 3건 (NCT ID, Phase)", "무료"],
    ],
    col_widths=[1.2, 3.5, 3.5, 6.3, 2]
)

h2(doc, "5.1 Orphanet 데이터 로컬 파싱")
body(doc, "Orphanet의 XML 파일들을 로컬에 다운로드하여 사용합니다. 외부 API 호출 없이 빠르게 조회할 수 있습니다.")
make_table(doc,
    ["파일명", "내용"],
    [
        ["en_product4.xml", "4,335개 희귀질환의 HPO 빈도 Annotation (매우 드묾/드묾/빈번/매우 빈번)"],
        ["en_product6.xml", "각 질환과 관련된 유전자 목록 + 유전자-질환 연관 유형 (Disease-causing 등)"],
        ["en_product9_ages.xml", "발병 연령(신생아/유아/성인 등) + 유전 양식 (상염색체 우성/열성 등)"],
        ["en_product9_prev.xml", "유병률 수치 (Point prevalence, 예: 1/100,000)"],
    ],
    col_widths=[5, 11.5]
)

h2(doc, "5.2 내부 DB 교차검증")
body(doc, "외부 API 결과를 받은 후, 로컬에 보유한 orphadata_weighted.csv (4,335개 질환, 115,878행)와 비교합니다.")
bullet(doc, "일치 시: 'DB·API 교차검증 일치' → 신뢰도 높음")
bullet(doc, "불일치 시: 'DB·API 불일치 — 추가 확인 필요' → 의사에게 주의 표시")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 6장 — AWS 리소스 구성
# ═══════════════════════════════════════════════════════════════════════

h1(doc, "6장. AWS 리소스 구성 목록")
body(doc, "이 프로젝트에서 사용 중인 AWS 리소스의 전체 목록입니다.")

h2(doc, "6.1 핵심 리소스 (현재 운영 중)")
make_table(doc,
    ["카테고리", "리소스 이름", "ID/ARN 요약", "상태"],
    [
        ["VPC", "say2-2team", "vpc-06dd0ad1f2335ea74 (10.0.0.0/24)", "✅ 운영"],
        ["S3", "say2-2team-bucket", "ap-northeast-2, SSE-KMS 암호화", "✅ 운영"],
        ["Lambda", "say2-2team-phase2-vision", "Python 3.11, 1GB RAM, 5분 타임아웃", "✅ 운영"],
        ["SageMaker", "say2-2team-soonet-endpoint", "ml.m5.large, PyTorch 2.0", "⏸ 발표 전 생성"],
        ["EC2", "2-2team-fhir-ec2", "i-0f3f223fd40217b12, t3.large, Ubuntu", "✅ 운영"],
        ["Aurora DB", "patient-db-cluster", "PostgreSQL 16.11, rarelinkai 스키마", "✅ 운영"],
        ["Cognito", "say2-2team-rare-link-pool", "ap-northeast-2_CMtZTRCTa, 5명", "✅ 운영"],
        ["CloudFront", "say2-2team-cf-distribution", "E2ZHONIV05TX9D, d300v14l8u0wx7.cloudfront.net", "✅ 배포 완료"],
        ["WAF", "say2-2team-waf", "us-east-1 글로벌, Essentials 보호", "✅ 운영"],
        ["KMS", "say2-2team-data-key", "4a1be264-1ccf-4a2c-a937-3e6847a751d5", "✅ 운영"],
        ["CloudTrail", "say2-2team-audit-trail", "S3 + CloudWatch 로그 저장", "✅ 운영"],
        ["SNS", "say2-2team-alerts", "Lambda 에러/SageMaker 지연/S3 삭제 알람", "✅ 운영"],
    ],
    col_widths=[3, 4.5, 6.5, 2.5]
)

h2(doc, "6.2 S3 버킷 구조")
code_block(doc,
"""say2-2team-bucket/
├── Phase_2/
│   ├── uploads/          ← 사용자가 업로드한 X-ray 이미지
│   ├── results/          ← AI 분석 결과 JSON
│   └── models/soonet/    ← model.tar.gz (SooNet 가중치 146MB)
├── frontend/             ← React 빌드 파일 (CloudFront가 서비스)
├── RAG/
│   └── rag_llm_3.py      ← EC2에서 실행하는 RAG 파이프라인 스크립트
├── final_reports/
│   └── {session_id}/report.pdf  ← 생성된 최종 소견서 PDF
├── cloudtrail-logs/      ← API 감사 로그
└── config-logs/          ← AWS Config 변경 이력"""
)

h2(doc, "6.3 비용 현황")
make_table(doc,
    ["리소스", "예상 비용", "비고"],
    [
        ["SageMaker Endpoint", "$0.115/시간", "⚠️ 발표 직전에만 켜고 즉시 삭제 필요"],
        ["EC2 (t3.large)", "~$0.075/시간", "RAG 실행 서버"],
        ["Aurora DB", "팀 공유 비용", "patient-db-cluster 공동 사용"],
        ["S3 저장", "$0.023/GB/월", "현재 약 200MB 사용"],
        ["CloudFront", "$0.0085/10,000건", "트래픽 기반"],
        ["WAF", "$5/WebACL/월", "고정 비용"],
        ["전체 알람 기준", "$80/일 초과 시", "SNS 이메일 알람 발송"],
    ],
    col_widths=[5, 4, 7.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 7장 — 데이터베이스 구조
# ═══════════════════════════════════════════════════════════════════════

h1(doc, "7장. 데이터베이스 구조 (Aurora PostgreSQL)")

h2(doc, "7.1 개요")
body(doc, "진단 과정의 모든 데이터는 Aurora PostgreSQL 클러스터(patient-db-cluster)의 rarelinkai 스키마에 저장됩니다. 팀 전체가 공유하는 리소스입니다.")
make_table(doc,
    ["항목", "값"],
    [
        ["Cluster 이름", "patient-db-cluster"],
        ["Engine", "Aurora PostgreSQL 16.11"],
        ["Writer Endpoint", "patient-db-cluster.cluster-cxmiyawwwhbt.ap-northeast-2.rds.amazonaws.com"],
        ["스키마", "rarelinkai"],
        ["접속 포트", "5432"],
        ["앱 사용자", "app_user (비밀번호: Secrets Manager 저장)"],
    ],
    col_widths=[4.5, 12]
)

h2(doc, "7.2 테이블 구조 및 INSERT 순서")
body(doc, "테이블 간에 외래키(FK) 관계가 있어서 반드시 아래 순서로 데이터를 저장해야 합니다.")
code_block(doc,
"""INSERT 순서:
  1. raw_emr_bundle       ← 원본 EMR 데이터 저장
        ↓ (FK)
  2. patient_profile      ← 환자 기본 프로파일
        ↓ (FK)
  3. diagnosis_session    ← 진단 세션 (status: pending → completed)
        ↓ (FK)
  4. final_report         ← 최종 소견서 JSON 저장

  5. rag_api_cache        ← API 호출 결과 캐시 (중복 요청 방지)"""
)

h2(doc, "7.3 주요 테이블 설명")
make_table(doc,
    ["테이블", "역할", "주요 컬럼"],
    [
        ["raw_emr_bundle", "원본 전자의무기록 저장", "session_id, emr_json, created_at"],
        ["patient_profile", "환자 기본정보", "patient_id, age, sex, chief_complaint"],
        ["diagnosis_session", "진단 처리 세션 상태 관리", "session_id, status (pending/processing/completed/error)"],
        ["final_report", "AI가 생성한 최종 소견서 JSON", "session_id, recommendation, clinical_notes, created_at"],
        ["rag_api_cache", "외부 API 응답 캐시", "query_key, response_json, expires_at"],
    ],
    col_widths=[4, 4.5, 8]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 8장 — 보안 구성
# ═══════════════════════════════════════════════════════════════════════

h1(doc, "8장. 보안 구성")
body(doc, "환자 데이터를 다루는 의료 시스템이므로 5개 보안 레이어를 적용했습니다.")

h2(doc, "8.1 보안 5계층")
make_table(doc,
    ["레이어", "항목", "적용 내용"],
    [
        ["Layer 1\n네트워크", "VPC Private Subnet", "Lambda, SageMaker를 인터넷 노출 없는 Private Subnet에 배치"],
        ["", "Security Group", "Lambda → SageMaker 포트 443만 허용 (최소 권한)"],
        ["", "S3 퍼블릭 차단", "4개 항목 모두 활성화 (직접 URL 접근 불가)"],
        ["Layer 2\n암호화", "SSE-KMS", "S3 파일 저장 시 KMS 키로 암호화 (alias/say2-2team-data-key)"],
        ["", "CloudTrail 암호화", "감사 로그도 별도 KMS 키로 암호화"],
        ["Layer 3\n감사/탐지", "CloudTrail", "모든 AWS API 호출 + S3 데이터 이벤트 기록"],
        ["", "AWS Config", "리소스 변경 이력 추적"],
        ["Layer 4\n비밀 관리", "Secrets Manager", "DB 비밀번호, API 키 등을 코드에서 분리하여 안전 저장"],
        ["", "IAM 최소 권한", "Lambda는 Phase_2/* 폴더와 자기 Endpoint만 접근 가능"],
        ["Layer 5\n모니터링", "CloudWatch 알람", "Lambda 에러 ≥3회/5분, SageMaker 지연 ≥30초 시 이메일 발송"],
        ["", "비용 알람", "전체 $80/일, SageMaker $10/일 초과 시 알람"],
        ["추가\nWAF", "웹 방화벽", "DDoS, Layer7 공격, SQL Injection, XSS, Known bad inputs 차단"],
    ],
    col_widths=[2.5, 4, 10]
)

h2(doc, "8.2 Cognito 인증 흐름")
body(doc, "의사/사용자가 시스템에 로그인할 때 Amazon Cognito를 통해 인증합니다.")
code_block(doc,
"""사용자 → 로그인 (ID/PW) → Cognito User Pool (say2-2team-rare-link-pool)
    → JWT 토큰 발급 → API Gateway 요청 시 헤더에 포함
    → API Gateway가 Cognito에 토큰 검증 → Lambda 함수 실행 허용""")
bullet(doc, "현재 등록 사용자: 5명")
bullet(doc, "User Pool ID: ap-northeast-2_CMtZTRCTa")
bullet(doc, "App Client ID: 1280u1fg8gbvt1g21sv8dn4246")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 9장 — 배포 현황
# ═══════════════════════════════════════════════════════════════════════

h1(doc, "9장. 배포 현황 — CloudFront 사이트")

h2(doc, "9.1 배포 URL")
body(doc, "시스템은 AWS CloudFront를 통해 아래 URL로 배포되어 있습니다:")
code_block(doc, "https://d300v14l8u0wx7.cloudfront.net")
code_block(doc, "데모 모드: https://d300v14l8u0wx7.cloudfront.net/?demo=1")

h2(doc, "9.2 프론트엔드 기술 스택")
make_table(doc,
    ["기술", "역할"],
    [
        ["React 18", "UI 프레임워크"],
        ["Vite 5", "빌드 도구 (빠른 개발 서버)"],
        ["Tailwind CSS", "스타일링 (유틸리티 CSS)"],
        ["AWS CloudFront", "전 세계 CDN 배포"],
        ["AWS S3", "정적 파일 호스팅 (say2-2team-bucket/frontend/)"],
    ],
    col_widths=[4.5, 12]
)

h2(doc, "9.3 CloudFront 설정")
make_table(doc,
    ["항목", "설정값"],
    [
        ["Distribution ID", "E2ZHONIV05TX9D"],
        ["Origin", "say2-2team-bucket.s3.ap-northeast-2.amazonaws.com/frontend"],
        ["기본 파일", "index.html (SPA 라우팅)"],
        ["에러 처리", "403/404 에러 → index.html 반환 (200)"],
        ["캐시 전략", "index.html: no-cache | JS/CSS: 1년 immutable"],
        ["WAF 연결", "say2-2team-waf (us-east-1)"],
    ],
    col_widths=[4.5, 12]
)

h2(doc, "9.4 GitHub Actions 자동 배포")
body(doc, "코드 변경 시 자동으로 빌드 → S3 업로드 → CloudFront 캐시 무효화가 실행됩니다.")
code_block(doc,
"""# .github/workflows/deploy.yml
# main 브랜치에 push 시 자동 실행:
1. npm run build        ← React 앱 빌드
2. aws s3 sync dist/ s3://say2-2team-bucket/frontend/  ← S3 업로드
3. CloudFront 캐시 무효화  ← 사용자에게 최신 버전 제공""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 10장 — 검증 결과
# ═══════════════════════════════════════════════════════════════════════

h1(doc, "10장. 검증 결과")

h2(doc, "10.1 SooNet X-ray 모델 성능")
body(doc, "CheXpert 데이터셋 50장으로 성능 검증 결과입니다.")
make_table(doc,
    ["지표", "결과"],
    [
        ["평균 AUROC", "0.78"],
        ["흉막삼출 (Pleural Effusion)", "AUROC 0.93+"],
        ["경화 (Consolidation)", "AUROC 0.93+"],
        ["학습 데이터", "MIMIC-CXR (448×448 전처리)"],
    ],
    col_widths=[6, 10.5]
)
tip_box(doc, "AUROC 0.93은 93%의 정확도로 이상 소견을 감지한다는 의미입니다. 랜덤(0.5)보다 훨씬 높습니다.")

h2(doc, "10.2 LIRICAL 희귀질환 스코어링 성능 (참고)")
body(doc, "4월 27일 시점 검증 결과 (4,293개 전수 테스트):")
make_table(doc,
    ["지표", "결과", "의미"],
    [
        ["Recall@1", "81.6%", "정답 질환이 1등으로 나올 확률"],
        ["Recall@10", "98.3%", "정답 질환이 Top10 내에 있을 확률"],
        ["임상 시나리오 5개", "Recall@3 = 100%", "실제 케이스 5개 모두 Top3에 포함"],
        ["PMID 유효율", "100%", "AI가 만들어낸 가짜 논문 0건"],
    ],
    col_widths=[4, 3, 9.5]
)

h2(doc, "10.3 외부 API 검증 현황")
make_table(doc,
    ["API", "상태", "비고"],
    [
        ["AWS Bedrock (Claude Haiku)", "✅ 정상", "Phase 1 증상→HPO, Phase 4 랭킹 정리"],
        ["AWS Bedrock (Claude Sonnet 3.5)", "✅ 정상", "최종 소견서 생성"],
        ["PubMed E-utilities", "✅ 정상", "3건/질환, 초록 400자 이내"],
        ["ClinicalTrials.gov v2", "✅ 정상", "RECRUITING 상태만 필터"],
        ["PubCaseFinder", "✅ 정상", "target=omim, enrich 로직 포함"],
        ["Monarch Initiative", "✅ 정상", "인과 유전자 + OMIM 검색 폴백"],
        ["Orphanet (로컬 XML)", "✅ 정상", "4개 XML 파일 로컬 파싱"],
    ],
    col_widths=[5.5, 2, 9]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 11장 — 향후 계획
# ═══════════════════════════════════════════════════════════════════════

h1(doc, "11장. 향후 계획")

h2(doc, "11.1 5월 이후 남은 검증 작업")
make_table(doc,
    ["항목", "목표"],
    [
        ["MIMIC 실환자 재검증 (3명+)", "5단계 파이프라인 기준 전체 재실행"],
        ["RAGAS 평가", "Faithfulness ≥ 0.8, Answer Relevancy 측정"],
        ["AWS Bedrock 비용 측정", "Claude Sonnet 3.5 vs 4.6 비용 비교"],
        ["5개 API 병렬 호출 latency", "목표: 60초 이내"],
        ["MRN 누출 여부 검증", "summary 필드에 환자 식별 정보 포함 여부 확인"],
    ],
    col_widths=[5.5, 11]
)

h2(doc, "11.2 검토 중인 개선 사항")
make_table(doc,
    ["항목", "현재", "검토 방향"],
    [
        ["LLM 모델", "Claude Sonnet 3.5", "Claude Sonnet 4.6 업그레이드 검토"],
        ["RAG 수집 건수", "Top 3 × 각 3건", "Top 5~10으로 확대 검토"],
        ["내부 DB 교차검증 범위", "Top 3", "Top 10으로 확대 검토"],
        ["병렬 호출 방식", "asyncio", "rate limit 시 부분 순차 처리 검토"],
        ["역학 정보", "희귀질환만", "일반 질환 포함 검토"],
        ["Step Functions", "미구현", "진단 파이프라인 오케스트레이션 도입 예정"],
        ["ElastiCache Redis", "미구현", "HPO 버퍼 캐싱으로 응답 속도 향상 예정"],
    ],
    col_widths=[4, 4, 8.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 부록 — 용어 사전
# ═══════════════════════════════════════════════════════════════════════

h1(doc, "부록. 용어 사전")
body(doc, "이 보고서에서 자주 등장하는 전문 용어를 초보자를 위해 쉽게 설명합니다.")

make_table(doc,
    ["용어", "설명"],
    [
        ["RAG (Retrieval-Augmented Generation)", "AI가 답변 생성 전에 외부 데이터베이스를 실시간 검색하여 근거 있는 답변을 생성하는 기술"],
        ["HPO (Human Phenotype Ontology)", "사람의 증상을 전 세계 공통 코드로 표현하는 표준 체계 (예: HP:0012735 = 기침)"],
        ["Orphanet", "유럽에서 운영하는 희귀질환 데이터베이스. 전 세계 6,000+ 희귀질환 정보 보유"],
        ["OrphaCode", "Orphanet이 부여한 희귀질환 고유 번호 (예: ORPHA:586 = 낭성섬유증)"],
        ["LIRICAL", "HPO 기반 희귀질환 확률 계산 알고리즘 (Likelihood Ratio 방식)"],
        ["SooNet", "이 프로젝트에서 개발한 흉부 X-ray 분석 AI 모델 (DenseNet-121 + U-Net)"],
        ["AUROC", "AI 모델 성능 지표. 0.5 = 랜덤, 1.0 = 완벽. 0.8+ 이면 좋은 성능"],
        ["AWS Bedrock", "AWS에서 제공하는 AI 모델 호출 서비스. Claude, Titan 등 다양한 모델 사용 가능"],
        ["Lambda", "서버 없이 코드를 실행하는 AWS 서버리스 서비스. 요청이 들어올 때만 실행되어 비용 절감"],
        ["SageMaker", "AI/ML 모델을 학습하고 배포하는 AWS 서비스"],
        ["CloudFront", "전 세계 엣지 서버를 통해 웹사이트를 빠르게 서비스하는 AWS CDN"],
        ["Aurora PostgreSQL", "AWS 관리형 고성능 PostgreSQL 데이터베이스 서비스"],
        ["VPC (Virtual Private Cloud)", "AWS 내의 가상 사설 네트워크. 외부와 격리된 보안 구역 생성 가능"],
        ["KMS (Key Management Service)", "암호화 키를 안전하게 관리하는 AWS 서비스"],
        ["MIMIC-CXR", "MIT에서 공개한 흉부 X-ray 공개 데이터셋 (22만+ 환자)"],
        ["PubMed", "미국 NCBI가 운영하는 의학 논문 검색 데이터베이스 (3천만+ 논문)"],
        ["MDT (Multi-Disciplinary Team)", "다학제 팀 협진. 여러 전문과 의사가 함께 진단·치료 계획 수립"],
        ["FHIR (Fast Healthcare Interoperability Resources)", "의료 데이터 교환 국제 표준 형식"],
        ["SSE-KMS", "S3 파일을 KMS 키로 서버 측 암호화하는 방식 (Server-Side Encryption)"],
        ["WAF (Web Application Firewall)", "웹 방화벽. DDoS, SQL Injection, XSS 등 웹 공격을 자동 차단"],
    ],
    col_widths=[5.5, 11]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("─────────────────────────────────────────────")
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SKKU AWS SAY 2기 2팀  |  Rare-Link AI  |  2026-05-24")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

# ─────────────────────────── 저장 ─────────────────────────────────────

out = "/Users/skku_mac08/Library/CloudStorage/GoogleDrive-ihtwandy0528@gmail.com/내 드라이브/ aws_바이오헬스케어_성균관대학교/aws_say2_project/aws_say2_project_vision/note_정리/rag/RAG_구현_보고서_초보자용.docx"
doc.save(out)
print(f"저장 완료: {out}")
